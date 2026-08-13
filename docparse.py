"""Extract plain text from uploaded documents: PDF, DOCX, TXT, MD."""
import io
import re

try:
    from pypdf import PdfReader
except Exception:  # pragma: no cover
    PdfReader = None
try:
    import docx  # python-docx
except Exception:  # pragma: no cover
    docx = None


def _from_pdf(data: bytes) -> str:
    if PdfReader is None:
        raise RuntimeError("pypdf is not installed; cannot read PDFs")
    reader = PdfReader(io.BytesIO(data))
    pages = []
    for page in reader.pages:
        try:
            pages.append(page.extract_text() or "")
        except Exception:
            pages.append("")
    return "\n\n".join(pages)


def _from_docx(data: bytes) -> str:
    if docx is None:
        raise RuntimeError("python-docx is not installed; cannot read DOCX files")
    document = docx.Document(io.BytesIO(data))
    blocks = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                blocks.append(" | ".join(cells))
    return "\n\n".join(blocks)


EXTENSIONS = {
    ".pdf": "pdf",
    ".docx": "docx",
    ".doc": "docx",
    ".txt": "text",
    ".md": "markdown",
    ".markdown": "markdown",
}


def detect_type(filename: str) -> str:
    ext = "." + filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    return EXTENSIONS.get(ext, "text")


def extract_text(filename: str, data: bytes) -> str:
    """Return plain text for a file, given name + raw bytes."""
    kind = detect_type(filename)
    if kind == "pdf":
        text = _from_pdf(data)
    elif kind == "docx":
        text = _from_docx(data)
    else:
        try:
            text = data.decode("utf-8-sig")
        except UnicodeDecodeError:
            try:
                text = data.decode("latin-1")
            except Exception:
                text = data.decode("utf-8", errors="replace")
    # Normalise whitespace a little but keep paragraph breaks.
    text = re.sub(r"\r\n", "\n", text)
    text = re.sub(r"[ \t]+\n", "\n", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _split_long_paragraph(para: str, chunk_size: int, overlap: int) -> list[str]:
    """Split a single long paragraph into overlapping sub-chunks."""
    step = max(1, chunk_size - overlap)
    sub_chunks = []
    start = 0
    while start < len(para):
        end = min(start + chunk_size, len(para))
        chunk = para[start:end].strip()
        if chunk:
            sub_chunks.append(chunk)
        if end >= len(para):
            break
        start += step
    return sub_chunks


def chunk_text(text: str, chunk_size: int = 900, overlap: int = 150) -> list[str]:
    """Split text into overlapping chunks on paragraph boundaries.

    Falls back to windowed character splits with overlap for very long paragraphs.
    """
    text = text.strip()
    if not text:
        return []

    raw_paragraphs = [p.strip() for p in re.split(r"\n\s*\n", text) if p.strip()]
    if not raw_paragraphs:
        return []

    # Flatten any oversized paragraphs into sub-chunks first
    paragraphs: list[str] = []
    for p in raw_paragraphs:
        if len(p) > chunk_size:
            paragraphs.extend(_split_long_paragraph(p, chunk_size, overlap))
        else:
            paragraphs.append(p)

    chunks: list[str] = []
    current = ""
    for para in paragraphs:
        if not current:
            current = para
        elif len(current) + 2 + len(para) <= chunk_size:
            current += "\n\n" + para
        else:
            chunks.append(current)
            # Retain the trailing context of the previous chunk as overlap
            if overlap > 0 and len(current) > overlap:
                tail = current[-overlap:].strip()
                current = f"{tail}\n\n{para}" if tail else para
            else:
                current = para

    if current:
        chunks.append(current)
    return chunks